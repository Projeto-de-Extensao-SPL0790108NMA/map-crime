import csv
import io
from datetime import datetime, time
from typing import ClassVar

from django.db.models import Count
from django.http import HttpResponse
from django.utils import timezone
from django.utils.dateparse import parse_date, parse_datetime
from docx import Document
from drf_yasg import openapi
from drf_yasg.utils import swagger_auto_schema
from openpyxl import Workbook
from rest_framework.exceptions import ValidationError
from rest_framework.permissions import IsAuthenticated
from rest_framework.views import APIView

from accounts.permissions.groups import IsAdmin, IsUser
from api import choice
from api.models import Denuncia


class DenunciaReportView(APIView):
    """
    Gera relatórios de denúncias nos formatos CSV, XLSX e DOCX
    com resumo por status e listagem detalhada filtrada por período.
    """

    permission_classes: ClassVar = [IsAuthenticated, IsAdmin | IsUser]

    formato_param = openapi.Parameter(
        "formato",
        openapi.IN_QUERY,
        description="Formato de saída do relatório: csv, xlsx ou docs.",
        type=openapi.TYPE_STRING,
        required=False,
    )
    data_inicio_param = openapi.Parameter(
        "data_inicio",
        openapi.IN_QUERY,
        description="Data inicial (ISO-8601). Aceita YYYY-MM-DD ou YYYY-MM-DDThh:mm:ss (hora será ignorada).",
        type=openapi.TYPE_STRING,
        required=False,
    )
    data_fim_param = openapi.Parameter(
        "data_fim",
        openapi.IN_QUERY,
        description="Data final (ISO-8601). Aceita YYYY-MM-DD ou YYYY-MM-DDThh:mm:ss (hora será ignorada).",
        type=openapi.TYPE_STRING,
        required=False,
    )

    @swagger_auto_schema(
        tags=["Relatórios"],
        operation_description="Gera um relatório detalhado de denúncias com resumo por status.",
        responses={200: openapi.Response("Arquivo gerado no formato solicitado.")},
        manual_parameters=[formato_param, data_inicio_param, data_fim_param],
    )
    def get(self, request, *args, **kwargs):
        formato = self._normalize_format(request.query_params.get("formato", "csv"))

        data_inicio = request.query_params.get("data_inicio")
        data_fim = request.query_params.get("data_fim")

        start_dt = self._parse_datetime_bound(data_inicio, "data_inicio", is_start=True)
        end_dt = self._parse_datetime_bound(data_fim, "data_fim", is_start=False)

        if start_dt and end_dt and start_dt > end_dt:
            raise ValidationError({"detail": "data_inicio não pode ser maior que data_fim."})

        queryset = Denuncia.objects.all().order_by("created_at")
        if start_dt:
            queryset = queryset.filter(created_at__gte=start_dt)
        if end_dt:
            queryset = queryset.filter(created_at__lte=end_dt)

        summary = self._build_summary(queryset, start_dt, end_dt)
        details = self._build_details(queryset)

        if formato == "csv":
            content = self._build_csv(summary, details)
            content_type = "text/csv; charset=utf-8"
            extension = "csv"
        elif formato == "xlsx":
            content = self._build_xlsx(summary, details)
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            extension = "xlsx"
        else:  # docx
            content = self._build_docx(summary, details)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            extension = "docx"

        filename = f"relatorio_denuncias_{timezone.now().strftime('%Y%m%d_%H%M%S')}.{extension}"
        response = HttpResponse(content, content_type=content_type)
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    def _normalize_format(self, value: str) -> str:
        normalized = (value or "").lower()
        if normalized in ("doc", "docs", "docx"):
            return "docx"
        if normalized in ("csv", "xlsx"):
            return normalized
        if normalized == "":
            return "csv"
        raise ValidationError({"formato": "Formato inválido. Use csv, xlsx ou docs."})

    def _parse_datetime_bound(
        self, raw_value: str | None, field_name: str, *, is_start: bool
    ) -> datetime | None:
        if not raw_value:
            return None

        parsed_dt = parse_datetime(raw_value)
        if parsed_dt:
            parsed_date = parsed_dt.date()
        else:
            parsed_date = parse_date(raw_value)

        if parsed_date:
            fallback_time = time.min if is_start else time.max
            combined = datetime.combine(parsed_date, fallback_time)
            return self._ensure_timezone(combined)

        raise ValidationError(
            {field_name: "Formato inválido. Use YYYY-MM-DD ou YYYY-MM-DDThh:mm:ss."}
        )

    def _ensure_timezone(self, value: datetime) -> datetime:
        if timezone.is_naive(value):
            return timezone.make_aware(value, timezone.get_current_timezone())
        return timezone.localtime(value)

    def _build_summary(self, queryset, start_dt, end_dt):
        status_counts = self._status_counts_for_queryset(queryset)
        has_filters = any([start_dt, end_dt])
        global_queryset = Denuncia.objects.all()
        global_status_counts = status_counts if not has_filters else self._status_counts_for_queryset(global_queryset)

        return {
            "total": queryset.count(),
            "status_counts": status_counts,
            "total_geral": global_queryset.count(),
            "status_counts_geral": global_status_counts,
            "periodo": self._format_period(start_dt, end_dt),
            "has_filters": has_filters,
        }

    def _status_counts_for_queryset(self, queryset):
        status_choices = getattr(choice.STATUS_CHOICES, "choices", choice.STATUS_CHOICES)
        label_map = dict(status_choices)
        qs = queryset.order_by()  # remove ordenação para agregações corretas
        aggregated = {
            row["status"]: row["total"]
            for row in qs.values("status").annotate(total=Count("status"))
        }

        status_counts = []
        seen = set()
        for key, label in status_choices:
            status_counts.append({"status": key, "label": label, "total": aggregated.get(key, 0)})
            seen.add(key)

        for key, total in aggregated.items():
            if key not in seen:
                status_counts.append({"status": key, "label": label_map.get(key, key), "total": total})

        return status_counts

    def _build_details(self, queryset):
        status_choices = getattr(choice.STATUS_CHOICES, "choices", choice.STATUS_CHOICES)
        label_map = dict(status_choices)
        data = []
        for denuncia in queryset:
            created = timezone.localtime(denuncia.created_at)
            data.append(
                {
                    "protocolo": denuncia.protocolo,
                    "categoria": denuncia.categoria,
                    "status": denuncia.status,
                    "status_label": label_map.get(denuncia.status, denuncia.status),
                    "created_at": created,
                    "descricao": denuncia.descricao,
                }
            )
        return data

    def _format_period(self, start_dt, end_dt):
        if not start_dt and not end_dt:
            return "Período completo"

        fmt = "%d/%m/%Y %H:%M"
        parts = []
        if start_dt:
            parts.append(f"início: {timezone.localtime(start_dt).strftime(fmt)}")
        if end_dt:
            parts.append(f"fim: {timezone.localtime(end_dt).strftime(fmt)}")
        return " | ".join(parts)

    def _build_csv(self, summary, details):
        buffer = io.StringIO()
        writer = csv.writer(buffer, delimiter=";")

        writer.writerow(["Relatório de Denúncias"])
        writer.writerow([summary["periodo"]])
        writer.writerow([f"Total de denúncias: {summary['total']}"])
        writer.writerow([])
        writer.writerow(["Totais por status (intervalo)"])
        for row in summary["status_counts"]:
            writer.writerow([row["label"], row["total"]])

        if summary["has_filters"]:
            writer.writerow([])
            writer.writerow(["Totais por status (geral)"])
            for row in summary["status_counts_geral"]:
                writer.writerow([row["label"], row["total"]])

        writer.writerow([])
        writer.writerow(["Protocolo", "Categoria", "Status", "Data de criação", "Descrição"])
        for item in details:
            writer.writerow(
                [
                    item["protocolo"],
                    item["categoria"],
                    item["status_label"],
                    item["created_at"].strftime("%d/%m/%Y %H:%M"),
                    (item["descricao"] or "").replace("\n", " "),
                ]
            )

        return buffer.getvalue().encode("utf-8")

    def _build_xlsx(self, summary, details):
        wb = Workbook()
        ws = wb.active
        ws.title = "Relatório"

        row_idx = 1
        ws.cell(row=row_idx, column=1, value="Relatório de Denúncias")
        row_idx += 1
        ws.cell(row=row_idx, column=1, value=summary["periodo"])
        row_idx += 1
        ws.cell(row=row_idx, column=1, value=f"Total de denúncias: {summary['total']}")
        row_idx += 2

        ws.cell(row=row_idx, column=1, value="Totais por status (intervalo)")
        row_idx += 1
        for status in summary["status_counts"]:
            ws.cell(row=row_idx, column=1, value=status["label"])
            ws.cell(row=row_idx, column=2, value=status["total"])
            row_idx += 1

        if summary["has_filters"]:
            row_idx += 1
            ws.cell(row=row_idx, column=1, value="Totais por status (geral)")
            row_idx += 1
            for status in summary["status_counts_geral"]:
                ws.cell(row=row_idx, column=1, value=status["label"])
                ws.cell(row=row_idx, column=2, value=status["total"])
                row_idx += 1

        row_idx += 1
        headers = ["Protocolo", "Categoria", "Status", "Data de criação", "Descrição"]
        for col, header in enumerate(headers, start=1):
            ws.cell(row=row_idx, column=col, value=header)
        row_idx += 1

        for item in details:
            ws.cell(row=row_idx, column=1, value=item["protocolo"])
            ws.cell(row=row_idx, column=2, value=item["categoria"])
            ws.cell(row=row_idx, column=3, value=item["status_label"])
            ws.cell(row=row_idx, column=4, value=item["created_at"].strftime("%d/%m/%Y %H:%M"))
            ws.cell(row=row_idx, column=5, value=item["descricao"])
            row_idx += 1

        stream = io.BytesIO()
        wb.save(stream)
        return stream.getvalue()

    def _build_docx(self, summary, details):
        document = Document()
        document.add_heading("Relatório de Denúncias", level=1)
        document.add_paragraph(summary["periodo"])
        document.add_paragraph(f"Total de denúncias: {summary['total']}")

        document.add_heading("Totais por status (intervalo)", level=2)
        for row in summary["status_counts"]:
            document.add_paragraph(f"{row['label']}: {row['total']}", style="List Bullet")

        if summary["has_filters"]:
            document.add_heading("Totais por status (geral)", level=2)
            for row in summary["status_counts_geral"]:
                document.add_paragraph(f"{row['label']}: {row['total']}", style="List Bullet")

        document.add_heading("Detalhes das denúncias", level=2)
        table = document.add_table(rows=1, cols=5)
        headers = ["Protocolo", "Categoria", "Status", "Data de criação", "Descrição"]
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = header

        for item in details:
            cells = table.add_row().cells
            cells[0].text = item["protocolo"]
            cells[1].text = item["categoria"]
            cells[2].text = item["status_label"]
            cells[3].text = item["created_at"].strftime("%d/%m/%Y %H:%M")
            cells[4].text = item["descricao"] or ""

        stream = io.BytesIO()
        document.save(stream)
        return stream.getvalue()
